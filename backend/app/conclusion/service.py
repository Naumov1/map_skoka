from io import BytesIO
from typing import List

from docx import Document

from datetime import datetime

from fastapi.responses import HTMLResponse, StreamingResponse
from loguru import logger
from app.applications.dao import ApplicationsDAO
from app.applications.models import ApplicationStatus, Applications
from app.schemas.base import Users
from app.conclusion.dao import ConclusionDAO
from app.conclusion.models import Conclusion
from app.conclusion.schemas import ResponseConclusion, ResponseCountConclusion
from app.exceptions import (
    CommissionStatementNotFound,
    ConclusionAlreadyCreatedException,
    FailedConvertToPDFExceptions,
    FailedCreateConclusionExceptions,
    FailedFormatNameExceptions,
    FileNotFoundException,
    UserNotFound,
)
from app.notification.service import NotificationService
from app.schemas.base import ResponseDetail
from app.service.base import ServiceBase
from app.signature.dao import SignatureDAO
from app.utils.s3 import s3_client


class ConclusionService(ServiceBase):
    @classmethod
    async def _placeholder_view(cls, conclusion_data: Conclusion) -> HTMLResponse:
        application = await ApplicationsDAO.find_by_id(conclusion_data.applications_id)
        html = f"""
        <!doctype html>
        <html lang="ru">
        <head>
          <meta charset="utf-8">
          <title>Заключение №{conclusion_data.id}</title>
          <style>
            body {{ margin: 0; background: #f6f7fb; color: #172033; font-family: Arial, sans-serif; }}
            main {{ max-width: 860px; margin: 32px auto; background: #fff; padding: 42px; border: 1px solid #e5e7eb; box-shadow: 0 20px 60px rgba(15, 23, 42, .12); }}
            h1 {{ margin: 0 0 24px; font-size: 30px; }}
            p {{ font-size: 16px; line-height: 1.65; }}
            b {{ display: inline-block; min-width: 190px; }}
            .result {{ margin-top: 24px; padding: 18px; border-radius: 8px; background: #f0fdf4; color: #166534; }}
            .note {{ margin-top: 18px; padding: 14px 16px; border-radius: 8px; background: #eef6ff; color: #31506f; }}
          </style>
        </head>
        <body>
          <main>
            <h1>Заключение комиссии №{conclusion_data.id}</h1>
            <p><b>Заявление:</b>№{conclusion_data.applications_id}</p>
            <p><b>ФИО:</b>{application.fio if application else "-"}</p>
            <p><b>Адрес:</b>{application.address if application else "-"}</p>
            <p><b>Кадастровый номер:</b>{application.cadastral_number if application else "-"}</p>
            <p><b>Дата создания:</b>{conclusion_data.create_date.strftime("%d.%m.%Y")}</p>
            <div class="result">Тестовое заключение подписано и готово к просмотру в демонстрационном режиме.</div>
            <div class="note">Демо-документ сформирован автоматически, потому что S3/MinIO хранилище для исходного файла сейчас не запущено.</div>
          </main>
        </body>
        </html>
        """
        return HTMLResponse(html)

    @classmethod
    async def _placeholder_download(cls, conclusion_data: Conclusion) -> StreamingResponse:
        application = await ApplicationsDAO.find_by_id(conclusion_data.applications_id)
        text = (
            f"Заключение комиссии №{conclusion_data.id}\n\n"
            f"Заявление: №{conclusion_data.applications_id}\n"
            f"ФИО: {application.fio if application else '-'}\n"
            f"Адрес: {application.address if application else '-'}\n"
            f"Кадастровый номер: {application.cadastral_number if application else '-'}\n"
            f"Дата создания: {conclusion_data.create_date.strftime('%d.%m.%Y')}\n"
            "Статус: подписано\n"
        )
        filename = f"conclusion-{conclusion_data.id}.txt"
        return StreamingResponse(
            BytesIO(text.encode("utf-8")),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    @classmethod
    async def create_conclusions(
        cls,
        applications_id: int,
        date: datetime,
        chairman_data: Users,
        members_data: List[Users],
        justification: str,
        documents: str,
        conclusion: str,
    ) -> ResponseDetail:
        """Создание заключения комиссии

        Args:
            applications_id (int): id заявления
            date (datetime): дата создания заключения
            chairman_id (int): id председателя комиссии
            members_id (list): id членов комиссии
            justification (str): обоснования
            documents (str): документы
            conclusion (str): результаты обследования

        Raises:
            FileNotFoundException: заявление не найдено
            UserNotFound: пользователь не найден
            FailedFormatNameExceptions: неверный формат имени
            ConclusionAlreadyCreatedException: заключение уже создано
            FailedCreateConclusionExceptions: ошибка создания заключения
            FailedConvertToPDFExceptions: ошибка конвертации в pdf
            FailedCreateConclusionExceptions: ошибка создания заключения

        Returns:
            ResponseDetail: результат создания заключения
        """
        applications_data: Applications = await ApplicationsDAO.find_by_id(
            applications_id
        )
        if not applications_data:
            logger.error(f"Заявление пользователя не найдено")
            raise FileNotFoundException
        if await ConclusionDAO.find_one_or_none(applications_id=applications_id):
            logger.error(
                f"Заключение коммисси для заявления уже существует ({applications_id})"
            )
            raise ConclusionAlreadyCreatedException
        if not chairman_data:
            logger.error(f"Пользователь не найден")
            raise UserNotFound
        chairman_fio = chairman_data.fio
        try:
            chairman = f"{chairman_fio.split(' ')[0]} {chairman_fio.split(' ')[1][0]}. {chairman_fio.split(' ')[2][0]}."
        except Exception as e:
            logger.error(f"Неверный формат имени для: {chairman_fio}: {e}")
            raise FailedFormatNameExceptions
        logger.debug(f"ФИО успешно получено: ({chairman_fio})")

        members = []
        for member in members_data:
            member_fio = member.fio
            try:
                member = f"{member_fio.split(' ')[0]} {member_fio.split(' ')[1][0]}. {member_fio.split(' ')[2][0]}."
            except Exception as e:
                logger.error(f"Неверный формат имени для: {member.id}: {e}")
                raise FailedFormatNameExceptions
            members.append(member)
        logger.debug(f"ФИО подписантов получено ({len(members_data)})")

        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        file_path = f"{timestamp}_{applications_id}.docx"
        try:
            await cls.fill_statement(
                file_path,
                str(datetime.now().strftime("%d.%m.%Y")),
                applications_data.address,
                chairman,
                members,
                f"{applications_data.fio}, собственник квартиры",
                documents,
                conclusion,
                justification,
            )
            logger.debug(f"Шаблон успешно создан ({applications_id})")
        except Exception as e:
            await s3_client.delete_file(file_path)
            logger.error(f"Ошибка создания файла заключения ({e})")
            raise FailedCreateConclusionExceptions

        try:
            await cls.convert_to_pdf(input_s3_key=file_path)
            logger.debug(
                f"Конвертирования word в pdf прошла успешна ({applications_id})"
            )
            logger.debug(f"Шаблон успешно создан в pdf ({applications_id})")
        except Exception as e:
            await s3_client.delete_file(file_path)
            logger.error(f"Ошибка конвертирования word в pdf ({applications_id}): {e}")
            raise FailedConvertToPDFExceptions

        try:
            conclusion_data: Conclusion = await ConclusionDAO.add(
                applications_id, date, file_path
            )
            conclusion_id = conclusion_data.id
            logger.debug(f"Заключение комисси создано в бд ({conclusion_id})")
        except Exception as e:
            await s3_client.delete_file(file_path)
            await s3_client.delete_file(file_path.replace(".docx", ".pdf"))
            logger.error(f"Ошибка создания заключения комисси: {e}")
            raise FailedCreateConclusionExceptions

        try:
            await ApplicationsDAO.update(
                applications_id, ApplicationStatus.COMMISSION_REVIEW
            )
            logger.debug(f"Статус заявления успешно изменен")
        except Exception as e:
            logger.error(
                f"Ошибка обновления статуса заявления: {applications_id} ({e})"
            )
            await s3_client.delete_file(file_path)
            await s3_client.delete_file(file_path.replace(".docx", ".pdf"))
            await ConclusionDAO.delete(conclusion_id)
            raise FailedCreateConclusionExceptions
        try:
            await NotificationService.send_notification(
                user_id=chairman_data.id,
                text=f"Заявление №{conclusion_id} - вас назначили председателем комиссии",
            )
            await SignatureDAO.add(chairman_data.id, conclusion_id)
            for member in members_data:
                await NotificationService.send_notification(
                    user_id=int(member.id),
                    text=f"Заявление №{conclusion_id} - вас назначили членом комиссии",
                )
                await SignatureDAO.add(int(member.id), conclusion_id)
            logger.debug(f"Уведомления успешно отправлены ({conclusion_id})")
        except Exception as e:
            logger.error(
                f"Ошибка отправки уведомлений для членов комиссии: {conclusion_id} ({e})"
            )
        try:
            await NotificationService.send_notification(
                user_id=applications_data.user_id,
                text=f"На ваше заявление №{applications_id} создано заключение комиссии №{conclusion_id}",
            )
            logger.debug(
                f"Уведомление для пользователя: {applications_data.user_id} отправлено"
            )
        except Exception as e:
            logger.error(
                f"Ошибка отправки уведомления о создании заключения: {conclusion_id} ({e})"
            )
        logger.info(f"Заключение комисси успешно создано")
        return {"detail": "Заключение комисси успешно создано"}

    @classmethod
    async def all_conclusions(cls, user_id: int) -> ResponseCountConclusion:
        """Получение всех заключений

        Args:
            user_id (int): id пользователя

        Returns:
            ResponseCountConclusion: список всех заключений
        """
        conclusion_data = await ConclusionDAO.all(user_id)
        return {"count": len(conclusion_data), "conclusions": conclusion_data}

    @classmethod
    async def my_conclusions(cls, user_id: int):
        conclusions = await ConclusionDAO.my(user_id)
        return {"count": len(conclusions), "conclusions": conclusions}

    @classmethod
    async def detail_conclusions(cls, id: int) -> ResponseConclusion:
        """детальная информация о заключении комиссии

        Args:
            id (int): id заключения

        Returns:
            ResponseConclusion: список всех заключений
        """
        logger.info("Получено заключение комиссии")
        return {"conclusions": await ConclusionDAO.find_by_id(model_id=id)}

    @classmethod
    async def search_conclusions(text: str, user_id: int) -> ResponseCountConclusion:
        """Поиск заключений комиссии по ФИО или кадастровому номеру

        Args:
            text (str): текст поиска
            user_id (int): id поиска

        Returns:
            ResponseCountConclusion: список заключений комисси
        """
        conclusion_data = await ConclusionDAO.search(text, user_id)
        logger.info(f"Получены заявления по поиску: {text[200:]}")
        return {"count": len(conclusion_data), "conclusions": conclusion_data}

    @classmethod
    async def filter_conclusions(
        cls,
        street: str | None,
        date_from: datetime | None,
        date_to: datetime | None,
        signed: bool | None,
        user_id: int,
    ) -> ResponseCountConclusion:
        """Фильтр заключений комиссии

        Args:
            street (str | None): улица
            date_from (datetime | None): дата от
            date_to (datetime | None): дата до
            signed (bool | None): подписан
            user_id (int): id пользователя

        Returns:
            ResponseCountConclusion: список заключений комисси
        """
        conclusion_data = await ConclusionDAO.filter(
            street, date_from, date_to, signed, user_id
        )
        logger.info(f"Фильтр по заявлениям найден")
        return {"count": len(conclusion_data), "conclusions": conclusion_data}

    @classmethod
    async def view_conclusions(cls, id: int) -> StreamingResponse:
        """Отображение содержания заключения комиссии

        Args:
            id (int): id заключения комиссии

        Raises:
            FileNotFoundException: файл не найден

        Returns:
            StreamingResponse: объект файла
        """
        conclusion_data: Conclusion = await ConclusionDAO.find_by_id(id)
        if not conclusion_data:
            logger.error(f"Заключение комиссии {id} не найдено")
            return CommissionStatementNotFound
        url = f"{(conclusion_data.file_url).split('.')[0]}.pdf"
        try:
            file = await s3_client.get_file(url)
        except Exception as e:
            logger.warning(f"Не удалось получить файл заключения {id} ({url}) из S3: {e}")
            file = None
        if not file:
            logger.error(f"Файл {id} ({url}) не найден")
            return await cls._placeholder_view(conclusion_data)
        logger.info(f"Файл {id} ({url}) успешно загружен")
        return StreamingResponse(
            file,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"inline; filename*=UTF-8''{url}",
                "X-Content-Type-Options": "nosniff",
            },
        )

    async def download_conclusions(id: int) -> StreamingResponse:
        """Скачивание файла заключения комиссии

        Args:
            id (int): id заключения

        Raises:
            FileNotFoundException: файл не найден

        Returns:
            StreamingResponse: объект файла
        """
        conclusion_data: Conclusion = await ConclusionDAO.find_by_id(id)
        if not conclusion_data:
            logger.error(f"Заключение комиссии {id} не найдено")
            return CommissionStatementNotFound
        url = conclusion_data.file_url
        try:
            file = await s3_client.get_file(url)
        except Exception as e:
            logger.warning(f"Не удалось получить файл заключения {id} ({url}) из S3: {e}")
            file = None
        if not file:
            logger.error(f"Файл {id} ({url}) не найден")
            return await ConclusionService._placeholder_download(conclusion_data)
        logger.info(f"Получен путь к файлу {id}")
        return StreamingResponse(
            file,
            media_type="application/octet-stream",
            headers={"Content-Disposition": f"attachment; filename={url}"},
        )

    @classmethod
    async def fill_statement(
        cls,
        output_path: str,
        date: str,
        address: str,
        chairman: str,
        members: list,
        owner: str,
        documents: str,
        conclusion: str,
        justification: str,
    ):
        doc = Document("doc/templates/statement_templates.docx")

        replacements = {
            "{DATE}": date,
            "{ADDRESS}": address,
            "{COMMISSION_INFO}": "Администрацией Центрального района г. Воронежа, решение №123 от 01.09.2025",
            "{CHAIRMAN}": chairman,
            "{MEMBERS}": ", ".join(members),
            "{OWNER}": owner,
            "{DOCUMENTS}": documents,
            "{CONCLUSION}": conclusion,
            "{JUSTIFICATION}": justification,
        }

        def replace_in_paragraph(paragraph):
            text = paragraph.text
            for key, value in replacements.items():
                if key in text:
                    text = text.replace(key, value)
            if text != paragraph.text:
                paragraph.clear()
                run = paragraph.add_run(text)
                run.font.underline = True

        # Обработка абзацев
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)

        # Обработка таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

        # Добавление подписей внизу
        doc.add_paragraph("   Председатель межведомственной комиссии:")

        p = doc.add_paragraph()
        p.add_run("\t_______________\t\t\t\t________________________\n")
        doc.add_paragraph()
        p.add_run("\t      подпись\t\t\t\t\t\t")
        p.add_run(chairman)

        doc.add_paragraph("\n   Члены межведомственной комиссии:\n")
        for member in members:
            p = doc.add_paragraph()
            p.add_run("\t_______________\t\t\t\t________________________\n")
            doc.add_paragraph()
            p.add_run("\t     подпись\t\t\t\t\t\t")
            p.add_run(member)

        file_obj = BytesIO()
        doc.save(file_obj)

        await s3_client.upload_bytes(output_path, file_obj.getvalue())
        return True

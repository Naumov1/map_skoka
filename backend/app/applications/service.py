from datetime import datetime
from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt
from fastapi.responses import HTMLResponse, StreamingResponse
from loguru import logger

from app.applications.dao import ApplicationsDAO
from app.applications.models import ApplicationStatus, Applications
from app.applications.schemas import (
    CommissionAnalysis,
    ResponseApplication,
    ResponseCommissionAnalysis,
    ResponseCountApplications,
    ResponseApplications,
    ResponseDetailApplications,
    ResponseRiskMapApplications,
    ResponseStreetsApplications,
)
from app.exceptions import (
    ApplicationDeleteError,
    ApplicationNotFound,
    FailedCreateApplicationExceptions,
    FailedFormatNameExceptions,
    FailedSendNotificationExceptions,
    InvalidFormatDepartureException,
    PermissionDeniedException,
)
from app.notification.service import NotificationService
from app.schemas.base import ResponseDetail, Role
from app.service.base import ServiceBase
from app.utils.s3 import s3_client


class ApplicationsService(ServiceBase):
    @classmethod
    def _make_commission_analysis(cls, application: Applications) -> CommissionAnalysis:
        problem = (application.problem or "").strip()
        text = f"{problem} {application.address or ''}".lower()
        patterns = [
            (("фасад", "штукатур", "строительные материалы", "падают", "проезж", "оград"), "Аварийное состояние фасада", "Высокая", ["представитель ТСЖ или управляющей организации", "инженер по эксплуатации здания", "специалист по фасадным работам", "аварийная служба при угрозе падения элементов"], ["Немедленно оградить опасную территорию возле фасада и установить предупреждающие знаки.", "Провести осмотр фасада с фотофиксацией разрушенных и ослабленных участков.", "Составить акт комиссии о состоянии фасада и рисках для граждан.", "Организовать срочное удаление аварийных элементов, которые могут упасть.", "Назначить ремонт фасада с ответственным исполнителем и контрольным сроком."], ["травмирование граждан", "падение элементов фасада", "ответственность за ненадлежащее содержание дома"]),
            (("аварийн", "фундамент", "кирпич", "обруш", "несущ"), "Аварийное состояние дома", "Высокая", ["межведомственная комиссия", "инженер-строитель", "представитель управляющей компании", "администрация муниципалитета"], ["Назначить срочный осмотр несущих конструкций и фундамента.", "Зафиксировать угрозу жизни и состояние конструкций актом комиссии.", "При подтверждении риска подготовить предписание и материалы для дальнейшего обращения."], ["угроза жизни и здоровью", "возможное разрушение конструкций"]),
            (("температур", "горячей воды", "горячая вода", "градус", "перерасчет"), "Некачественная коммунальная услуга: горячая вода", "Средняя", ["представитель управляющей компании", "специалист ресурсоснабжающей организации", "сантехник или инженер по тепловому узлу", "жилищная инспекция при бездействии исполнителя"], ["Назначить комиссионный замер температуры горячей воды в квартире заявителя.", "Составить акт проверки с указанием даты, времени, температуры и точки водоразбора.", "Проверить тепловой узел, циркуляцию горячей воды и параметры подачи в доме.", "При подтверждении отклонения выполнить перерасчет платы за период оказания услуги ненадлежащего качества.", "Установить срок восстановления температуры горячей воды до нормативных значений."], ["ненадлежащее качество коммунальной услуги", "переплата за услугу", "жалоба в жилищную инспекцию"]),
            (("отоплен", "батар", "нет тепла", "холодно", "температура в квартире"), "Некачественная коммунальная услуга: отопление", "Высокая", ["представитель управляющей компании", "теплотехник", "ресурсоснабжающая организация", "аварийно-диспетчерская служба"], ["Провести замер температуры воздуха в помещении и температуры теплоносителя.", "Составить акт проверки качества отопления.", "Проверить стояки, радиаторы, запорную арматуру и параметры подачи тепла.", "Устранить причину недостаточного отопления и назначить контрольный повторный замер.", "При подтверждении нарушения подготовить перерасчет платы."], ["нарушение санитарных условий", "замерзание помещения", "жалоба в жилищный надзор"]),
            (("горяч", "водоснаб", "стояк", "сантех", "прорыв"), "Проблема с водоснабжением или сантехникой", "Высокая", ["аварийная служба", "сантехник управляющей компании", "представитель УК"], ["Организовать выезд аварийной службы и перекрытие проблемного участка при необходимости.", "Составить акт осмотра с указанием причины и зоны ответственности.", "Назначить ремонт или замену поврежденного элемента с контрольным сроком."], ["ущерб имуществу", "ограничение коммунальной услуги", "повторная авария"]),
            (("крыш", "кровл", "протеч", "затоп", "влага", "плесень"), "Протечка кровли или залив", "Высокая", ["представитель УК", "специалист по кровле", "комиссия по осмотру помещения"], ["Провести осмотр кровли, подъезда и пострадавших помещений.", "Составить акт протечки с фотофиксацией и причиной повреждений.", "Определить срочный ремонт и меры по предотвращению повторного залива."], ["повреждение отделки", "плесень и влажность", "дальнейшее затопление"]),
            (("лифт", "застр", "кабина", "подъемник"), "Неисправность лифта", "Высокая", ["представитель УК", "лифтовая обслуживающая организация", "аварийная служба лифта"], ["Проверить факт неисправности и наличие угрозы безопасности.", "Передать заявку в специализированную лифтовую организацию.", "При опасной работе остановить лифт до устранения неисправности.", "Зафиксировать сроки ремонта и результат проверки."], ["опасность для жильцов", "ограничение доступа в квартиры", "повторное застревание"]),
            (("свет", "освещ", "фонар", "электр", "ламп"), "Освещение мест общего пользования", "Средняя", ["электрик управляющей компании", "представитель УК"], ["Проверить электросеть и светильники в местах общего пользования.", "Восстановить освещение и зафиксировать выполненные работы.", "Проверить безопасность прохода жильцов в темное время суток."], ["травматизм жильцов", "нарушение содержания общего имущества"]),
            (("мусор", "гряз", "уборк", "контейнер", "снег", "двор"), "Ненадлежащее содержание территории или подъезда", "Средняя", ["представитель УК", "подрядчик по уборке", "дворник или обслуживающая организация"], ["Проверить состояние подъезда, двора или контейнерной площадки.", "Составить акт ненадлежащего содержания с фотофиксацией.", "Организовать уборку и установить график контроля."], ["санитарные нарушения", "повторные жалобы жильцов", "предписание надзорного органа"]),
            (("присво", "помещение общего пользования", "мусоропровод", "металлическ", "двер"), "Самовольное занятие общего имущества", "Средняя", ["представитель управляющей компании", "члены комиссии", "собственник или лицо, ограничившее доступ", "жилищная инспекция при отказе устранить нарушение"], ["Проверить помещение общего пользования и подтвердить факт ограничения доступа.", "Составить акт комиссии с фотофиксацией двери, замка и фактического использования помещения.", "Выдать требование демонтировать дверь и освободить помещение общего пользования.", "При отказе подготовить обращение в жилищную инспекцию, пожарный надзор или прокуратуру."], ["нарушение прав собственников", "ограничение доступа к общему имуществу", "пожарная безопасность"]),
            (("квитанц", "начисл", "переплат", "тариф", "плата"), "Спор по начислениям или перерасчету", "Низкая", ["бухгалтерия управляющей компании", "расчетный центр", "представитель УК"], ["Проверить начисления по лицевому счету и основания для расчета.", "Сверить показания приборов учета, тарифы и период начисления.", "Подготовить письменный ответ заявителю и выполнить перерасчет при подтверждении ошибки."], ["переплата", "спор с собственником", "жалоба в жилищную инспекцию"]),
            (("балкон", "окн", "подъезд", "двер", "ступен"), "Ремонт общего имущества", "Средняя", ["представитель УК", "инженер по эксплуатации", "ремонтная бригада"], ["Осмотреть поврежденные элементы и определить принадлежность к общему имуществу.", "Составить дефектную ведомость и срок устранения.", "Проконтролировать ремонт и повторный осмотр после выполнения работ."], ["травматизм", "дальнейшее разрушение", "жалобы жителей"]),
        ]
        category = "Обращение по содержанию жилья"
        urgency = "Средняя"
        who_needed = ["представитель управляющей компании", "члены комиссии"]
        actions = ["Изучить заявление и приложенные материалы.", "Назначить осмотр объекта с фотофиксацией.", "Составить акт комиссии и определить ответственного исполнителя."]
        risks = ["нарушение срока рассмотрения", "повторное обращение заявителя"]

        for keywords, found_category, found_urgency, found_who, found_actions, found_risks in patterns:
            if any(keyword in text for keyword in keywords):
                category, urgency, who_needed, actions, risks = found_category, found_urgency, found_who, found_actions, found_risks
                break

        if "сосед" in text or "компенсац" in text or "ущерб" in text:
            category = "Залив или ущерб от соседей"
            urgency = "Средняя"
            who_needed = ["пострадавший заявитель", "сосед или собственник помещения", "представитель УК", "оценщик ущерба"]
            actions = ["Составить акт о заливе с участием сторон и управляющей компании.", "Зафиксировать повреждения фото и указать предварительную причину.", "Рекомендовать досудебное урегулирование или подготовку материалов для взыскания ущерба."]
            risks = ["материальный ущерб", "спор о размере компенсации", "судебное разбирательство"]

        return CommissionAnalysis(
            applicant=application.fio,
            address=application.address,
            category=category,
            urgency=urgency,
            what_happened=problem or "Описание проблемы в заявлении не сохранено. Используйте документ заявления и данные объекта для ручной проверки.",
            who_needed=who_needed,
            recommended_actions=actions,
            commission_focus=["проверить фактическое состояние объекта", "определить ответственного за устранение проблемы", "зафиксировать сроки и результат решения"],
            risks=risks,
        )

    @classmethod
    async def analyze_for_commission(cls, id: int, user_id: int, role: Role) -> ResponseCommissionAnalysis:
        application = await ApplicationsDAO.find_by_id(id)
        await cls.check_access_applications(id, application, user_id, role)
        analysis = cls._make_commission_analysis(application)
        await ApplicationsDAO.save_commission_analysis(id, analysis.model_dump_json(ensure_ascii=False))
        return {"analysis": analysis}

    @classmethod
    def _placeholder_view(cls, application: Applications) -> HTMLResponse:
        status = getattr(application.status, "value", application.status)
        return HTMLResponse(f"""<!doctype html><html lang="ru"><head><meta charset="utf-8"><title>Заявление №{application.id}</title><style>body{{margin:0;background:#f6f7fb;color:#172033;font-family:Arial,sans-serif}}main{{max-width:860px;margin:32px auto;background:#fff;padding:42px;border:1px solid #e5e7eb;box-shadow:0 20px 60px rgba(15,23,42,.12)}}p{{font-size:16px;line-height:1.65}}b{{display:inline-block;min-width:190px}}.note{{margin-top:28px;padding:14px 16px;border-radius:8px;background:#eef6ff;color:#31506f}}</style></head><body><main><h1>Заявление №{application.id}</h1><p><b>ФИО:</b>{application.fio}</p><p><b>Телефон:</b>{application.phone}</p><p><b>Email:</b>{application.email}</p><p><b>Адрес:</b>{application.address}</p><p><b>Кадастровый номер:</b>{application.cadastral_number}</p><p><b>Описание:</b>{application.problem or 'Описание проблемы не сохранено'}</p><p><b>Статус:</b>{status}</p><div class="note">Демо-документ сформирован автоматически, потому что файл заявления недоступен в S3/MinIO.</div></main></body></html>""")

    @classmethod
    def _placeholder_download(cls, application: Applications) -> StreamingResponse:
        status = getattr(application.status, "value", application.status)
        text = f"Заявление №{application.id}\n\nФИО: {application.fio}\nТелефон: {application.phone}\nEmail: {application.email}\nАдрес: {application.address}\nКадастровый номер: {application.cadastral_number}\nОписание: {application.problem or 'Описание проблемы не сохранено'}\nСтатус: {status}\n"
        return StreamingResponse(BytesIO(text.encode("utf-8")), media_type="text/plain; charset=utf-8", headers={"Content-Disposition": f"attachment; filename=application-{application.id}.txt"})

    @classmethod
    async def check_access_applications(cls, id: int, application: Applications | None, user_id: int, role: Role):
        if not application:
            raise ApplicationNotFound
        if application.user_id != user_id and role == Role.USERS:
            raise PermissionDeniedException
        return True

    @classmethod
    async def all(cls):
        return {"applications": await ApplicationsDAO.find_all()}

    @classmethod
    async def create_applications(cls, user_id: int, fio: str, phone: str, email: str, cadastral_number: str, problem: str, address: str) -> ResponseDetailApplications:
        file_path = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{user_id}.docx"
        try:
            parts = fio.split()
            applicant = f"{parts[0]} {parts[1][0]}. {parts[2][0]}."
        except Exception as e:
            logger.error(f"Invalid fio format for {user_id}: {e}")
            raise FailedFormatNameExceptions

        try:
            street_match = re.search(r"ул\.\s[^,]+", address, flags=re.IGNORECASE)
            application_data = await ApplicationsDAO.add(user_id, fio, phone, email, cadastral_number, street_match.group() if street_match else address, address, problem, file_path)
        except Exception as e:
            logger.error(f"Failed create application ({user_id}): {e}")
            raise FailedCreateApplicationExceptions

        if not await cls.applicant_statement(file_path, applicant, email, phone, problem, address):
            logger.warning(f"Continue without application file ({user_id})")
        try:
            await cls.convert_to_pdf(input_s3_key=file_path)
        except Exception as e:
            logger.warning(f"Continue without application PDF ({user_id}): {e}")

        try:
            await NotificationService.send_notification(user_id=user_id, text=f"Ваше заявление №{application_data.id} принято")
        except Exception as e:
            logger.error(f"Failed send notification ({user_id}): {e}")
            raise FailedSendNotificationExceptions
        return {"detail": "Заявление успешно создано", "applications": application_data}

    @classmethod
    async def delete_applications(cls, id: int, user_id: int, role: Role) -> ResponseApplications:
        application_data = await ApplicationsDAO.find_by_id(id)
        await cls.check_access_applications(id, application_data, user_id, role)
        if application_data.status != ApplicationStatus.ACCEPTED:
            raise ApplicationDeleteError
        application_data = await ApplicationsDAO.delete(id)
        try:
            await s3_client.delete_file(application_data.file_url)
            await s3_client.delete_file(application_data.file_url.replace(".docx", ".pdf"))
        except Exception as e:
            logger.warning(f"Failed delete files: {e}")
        return {"detail": "Заявление успешно удалено", "applications": application_data}

    @classmethod
    async def all_applications(cls) -> ResponseCountApplications:
        data = await ApplicationsDAO.all()
        return {"count": len(data), "applications": data}

    @classmethod
    async def risk_map_applications(cls) -> ResponseRiskMapApplications:
        data = await ApplicationsDAO.all()
        applications = []
        for application in data:
            status = application.get("status")
            applications.append({
                "id": application.get("id"),
                "address": application.get("address"),
                "street": application.get("street"),
                "problem": application.get("problem"),
                "commission_analysis": application.get("commission_analysis"),
                "status": getattr(status, "value", status),
                "departure_date": application.get("departure_date"),
            })
        return {"count": len(applications), "applications": applications}

    @classmethod
    async def my_applications(cls, user_id: int):
        data = await ApplicationsDAO.my(user_id=user_id)
        return {"count": len(data), "applications": data}

    @classmethod
    async def detail_applications(cls, id: int, user_id: int, role: Role) -> ResponseApplication:
        application = await ApplicationsDAO.find_by_id(id)
        await cls.check_access_applications(id, application, user_id, role)
        return {"applications": application}

    @classmethod
    async def addres_api(cls, search: str | None = None) -> ResponseStreetsApplications:
        return {"streets": await ApplicationsDAO.get_street(search)}

    @classmethod
    async def download_applications(cls, id: int, user_id: int, role: Role) -> StreamingResponse:
        application = await ApplicationsDAO.find_by_id(id)
        await cls.check_access_applications(id, application, user_id, role)
        try:
            file = await s3_client.get_file(application.file_url)
        except Exception as e:
            logger.warning(f"Failed get file {id}: {e}")
            file = None
        if not file:
            return cls._placeholder_download(application)
        return StreamingResponse(file, media_type="application/octet-stream", headers={"Content-Disposition": f"attachment; filename={application.file_url}"})

    @classmethod
    async def view_applications(cls, id: int, user_id: int, role: Role) -> StreamingResponse:
        application = await ApplicationsDAO.find_by_id(id)
        await cls.check_access_applications(id, application, user_id, role)
        url = f"{application.file_url.split('.')[0]}.pdf"
        try:
            file = await s3_client.get_file(url)
        except Exception as e:
            logger.warning(f"Failed get file {id}: {e}")
            file = None
        if not file:
            return cls._placeholder_view(application)
        return StreamingResponse(file, media_type="application/pdf", headers={"Content-Disposition": f"inline; filename*=UTF-8''{url}", "X-Content-Type-Options": "nosniff"})

    @classmethod
    async def search_applications(cls, text: str) -> ResponseCountApplications:
        data = await ApplicationsDAO.search(text)
        return {"count": len(data), "applications": data}

    @classmethod
    async def filter_applications(cls, street: str | None, date_from: datetime | None, date_to: datetime | None, is_departure: bool | None) -> ResponseCountApplications:
        data = await ApplicationsDAO.filter(street, date_from, date_to, is_departure)
        return {"count": len(data), "applications": data}

    @classmethod
    async def update_departure(cls, applications_id: str, departure_date: datetime) -> ResponseDetail:
        if applications_id.count("-") == 1 and applications_id.count(",") <= 0:
            ids = list(range(int(applications_id.split("-")[0]), int(applications_id.split("-")[-1]) + 1))
        elif applications_id.count(",") > 0 and applications_id.count("-") == 0:
            ids = [int(j) for j in applications_id.split(",")]
        else:
            try:
                ids = [int(applications_id)]
            except Exception:
                raise InvalidFormatDepartureException
        applications_data = await ApplicationsDAO.departure(ids, departure_date)
        for application in applications_data:
            await NotificationService.send_notification(user_id=application.get("user_id"), text=f"Для заявки №{application.get('id')} назначен выезд на {departure_date.strftime('%d.%m.%Y')}")
        return {"detail": "Выезд успешно назначен"}

    @classmethod
    async def applicant_statement(cls, output_path, applicant: str, email: str, phone_number: str, problem: str, address: str) -> bool:
        try:
            template_path = Path("doc/templates/applications_templates.docx")
            if not template_path.exists():
                raise FileNotFoundError(f"Template not found: {template_path}")
            doc = Document(template_path)
            replacements = {"{DATE}": datetime.today().strftime("%d.%m.%Y"), "{APPLICANT}": applicant, "{EMAIL}": email, "{PHONE_NUMBER}": phone_number, "{PROBLEM}": problem, "{ADDRESS}": address}

            def replace_in_paragraph(paragraph):
                text = paragraph.text
                for key, value in replacements.items():
                    text = text.replace(key, value)
                if text != paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run(text)
                    run.font.underline = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)
            file_obj = BytesIO()
            doc.save(file_obj)
            await s3_client.upload_bytes(url=output_path, file=file_obj.getvalue())
            return True
        except Exception as e:
            logger.error(f"Application document processing error: {e}")
            return False
